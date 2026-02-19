"""Tests for document parsers."""

import os
from unittest.mock import patch

import pytest
from docx import Document

from pitchdeck.models import DocumentParseError
from pitchdeck.parsers import extract_document
from pitchdeck.parsers.docx_parser import extract_docx
from pitchdeck.parsers.pdf import extract_pdf


class TestPDFParser:
    def test_extract_pdf_success(self, tmp_path):
        pdf_path = str(tmp_path / "test.pdf")
        with patch("pitchdeck.parsers.pdf.pymupdf4llm") as mock_pymupdf:
            mock_pymupdf.to_markdown.return_value = "# Test Document\n\nContent here."
            # Create a dummy file so os.path.exists passes
            with open(pdf_path, "w") as f:
                f.write("dummy")
            result = extract_pdf(pdf_path)
            assert result == "# Test Document\n\nContent here."
            mock_pymupdf.to_markdown.assert_called_once_with(pdf_path)

    def test_extract_pdf_file_not_found(self):
        with pytest.raises(DocumentParseError, match="File not found"):
            extract_pdf("/nonexistent/file.pdf")

    def test_extract_pdf_parsing_error(self, tmp_path):
        pdf_path = str(tmp_path / "bad.pdf")
        with open(pdf_path, "w") as f:
            f.write("not a pdf")
        with patch("pitchdeck.parsers.pdf.pymupdf4llm") as mock_pymupdf:
            mock_pymupdf.to_markdown.side_effect = RuntimeError(
                "Invalid PDF"
            )
            with pytest.raises(DocumentParseError, match="Invalid PDF"):
                extract_pdf(pdf_path)


class TestDOCXParser:
    def test_extract_docx_success(self, tmp_path):
        docx_path = str(tmp_path / "test.docx")
        doc = Document()
        doc.add_heading("Title", level=1)
        doc.add_paragraph("Body paragraph text.")
        doc.add_heading("Subtitle", level=2)
        doc.add_paragraph("More content.")
        doc.save(docx_path)

        result = extract_docx(docx_path)
        assert "# Title" in result
        assert "## Subtitle" in result
        assert "Body paragraph text." in result

    def test_extract_docx_empty_doc(self, tmp_path):
        docx_path = str(tmp_path / "empty.docx")
        doc = Document()
        doc.save(docx_path)

        result = extract_docx(docx_path)
        assert result == ""

    def test_extract_docx_no_headings(self, tmp_path):
        docx_path = str(tmp_path / "noheadings.docx")
        doc = Document()
        doc.add_paragraph("First paragraph.")
        doc.add_paragraph("Second paragraph.")
        doc.save(docx_path)

        result = extract_docx(docx_path)
        assert "First paragraph." in result
        assert "Second paragraph." in result
        assert "#" not in result

    def test_extract_docx_file_not_found(self):
        with pytest.raises(DocumentParseError, match="File not found"):
            extract_docx("/nonexistent/file.docx")


class TestExtractDocument:
    def test_dispatches_to_pdf(self, tmp_path):
        pdf_path = str(tmp_path / "test.pdf")
        with open(pdf_path, "w") as f:
            f.write("dummy")
        with patch("pitchdeck.parsers.pdf.pymupdf4llm") as mock_pymupdf:
            mock_pymupdf.to_markdown.return_value = "PDF content"
            result = extract_document(pdf_path)
            assert result == "PDF content"

    def test_dispatches_to_docx(self, tmp_path):
        docx_path = str(tmp_path / "test.docx")
        doc = Document()
        doc.add_paragraph("DOCX content")
        doc.save(docx_path)

        result = extract_document(docx_path)
        assert "DOCX content" in result

    def test_unsupported_format(self):
        with pytest.raises(DocumentParseError, match="Unsupported format"):
            extract_document("test.txt")

    def test_case_insensitive_extension(self, tmp_path):
        pdf_path = str(tmp_path / "test.PDF")
        with open(pdf_path, "w") as f:
            f.write("dummy")
        with patch("pitchdeck.parsers.pdf.pymupdf4llm") as mock_pymupdf:
            mock_pymupdf.to_markdown.return_value = "Content"
            result = extract_document(pdf_path)
            assert result == "Content"
